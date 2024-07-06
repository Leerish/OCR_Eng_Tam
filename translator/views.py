# translator/views.py
import os 
from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from .forms import TranslationRequestForm
from .models import TranslationRequest
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pytesseract
from PIL import Image
from googletrans import Translator

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_from_image(image_path):
    try:
        img = Image.open(image_path)
    except FileNotFoundError:
        return "Image file not found."

    try:
        text = pytesseract.image_to_string(img, lang='tam')
    except pytesseract.TesseractError as e:
        return f"Error during OCR process: {e}"

    return text

def translate_text(text, dest_lang='en'):
    translator = Translator()
    try:
        translated_text = translator.translate(text, dest=dest_lang).text
    except Exception as e:
        return f"Error during translation process: {e}"

    return translated_text

def get_summary(text, sentences_count=2):
    sentences = text.split('.')
    summary = '. '.join(sentences[:sentences_count]).strip()
    return summary + '.' if not summary.endswith('.') else summary

def create_word_doc(image_path, newspaper_name, edition, translated_text, output_path):
    doc = Document()

    # Add newspaper name centered
    doc.add_heading(newspaper_name, level=1)

    # Add edition and date on the left near the margin
    current_date = datetime.now().strftime("%B %d, %Y")
    doc.add_paragraph(f"Edition: {edition}\nDate: {current_date}")

    # Add uploaded image
    doc.add_picture(image_path, width=Inches(3.0))

    # Add translated text with a line break
    doc.add_heading("Translated Text:", level=2)
    doc.add_paragraph(translated_text)

    # Save the Word document
    doc.save(output_path)

def translate(request):
    if request.method == 'POST':
        form = TranslationRequestForm(request.POST, request.FILES)
        if form.is_valid():
            translation_request = form.save()
            image_path = translation_request.image.path
            newspaper_name = translation_request.newspaper_name
            edition = translation_request.edition

            original_text = extract_text_from_image(image_path)
            if "Error" in original_text:
                return render(request, 'translator/translate.html', {'form': form, 'error': original_text})

            translated_text = translate_text(original_text)
            if "Error" in translated_text:
                return render(request, 'translator/translate.html', {'form': form, 'error': translated_text})

            summary = get_summary(translated_text, sentences_count=2)
            output_word_path = f"{image_path}.docx"
            create_word_doc(image_path, newspaper_name, edition, summary, output_word_path)

            return redirect('download', path=output_word_path)

    else:
        form = TranslationRequestForm()
    return render(request, 'translator/translate.html', {'form': form})

def download(request, path):
    file_path = path
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/msword")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response
    raise Http404
